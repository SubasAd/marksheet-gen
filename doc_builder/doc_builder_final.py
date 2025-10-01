import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL # <-- Import for vertical alignment

# Make sure your common_functions.py is in the same directory or a reachable path
from .common_functions import *

def _get_grade_from_marks(obtained_marks, full_marks):
    """
    Calculates a letter grade based on the percentage of marks.
    This is a helper function to fill in missing grade data for TH/IN parts.
    """
    if not isinstance(obtained_marks, (int, float)) or not isinstance(full_marks, (int, float)) or full_marks == 0:
        return "N/A"
    
    percentage = (obtained_marks / full_marks) * 100
    if percentage >= 90: return "A+"
    if percentage >= 80: return "A"
    if percentage >= 70: return "B+"
    if percentage >= 60: return "B"
    if percentage >= 50: return "C+"
    if percentage >= 40: return "C"
    if percentage >= 30: return "D+"
    if percentage >= 20: return "D"
    return "E"


def create_marksheet_docx(student_data, school_info, output_filename):
    """
    Generates a .docx marksheet with a doubled-row subject breakdown 
    (External/Theory and Internal rows under each subject).
    """
    doc = docx.Document()
    # Assuming this is for the final/annual term.
    terminal_key = "annual" 
    add_heading_component(student_data, terminal_key, school_info, doc)

    terminal_data = student_data.get(terminal_key, {})
    subjects = terminal_data.get('subjects', {})

    # --- MAIN TABLE ---
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    headers = ['S.N.', 'Subjects', 'Credit Hour', 'Marks Obtained', 'Grade Point', 'Grade', 'Final Grade','Remarks']
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header_text)
        run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    sn = 1
    grand_total_marks = 0

    for subject_name, details in subjects.items():
        external = details.get('external', {})
        internal = details.get('internal', {})

        # --- Add the First Row (Theory/External) ---
        row1_cells = table.add_row().cells
        row1_cells[0].text = str(sn)
        row1_cells[1].text = f"{subject_name.upper()} (TH)"
        row1_cells[2].text = str(external.get('credit_hour', ''))
        row1_cells[3].text = str(external.get('mark', ''))
        row1_cells[4].text = str(external.get('gp', ''))
        # Calculate grade for the theory part as it's not in the data
        th_grade = _get_grade_from_marks(external.get('mark'), external.get('full_marks'))
        row1_cells[5].text = th_grade

        # --- Add the Second Row (Internal) ---
        row2_cells = table.add_row().cells
        row2_cells[1].text = f"{subject_name.upper()} (IN)"
        row2_cells[2].text = str(internal.get('credit_hour', ''))
        row2_cells[3].text = str(internal.get('mark', ''))
        row2_cells[4].text = str(internal.get('gp', ''))
        # Calculate grade for the internal part
        in_grade = _get_grade_from_marks(internal.get('mark'), internal.get('full_marks'))
        row2_cells[5].text = in_grade
        
        # Add the final grade to the second row before merging
        row2_cells[6].text = str(details.get('final_grade', ''))

        # --- Merge the vertical cells ---
        row1_cells[0].merge(row2_cells[0])  # Merge S.N.
        row1_cells[6].merge(row2_cells[6])  # Merge Final Grade

        # --- Vertically center the content in merged cells ---
        for cell_index in [0, 6]:
            row1_cells[cell_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row1_cells[cell_index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Center align other relevant cells
        for r_cells in [row1_cells, row2_cells]:
            for i in range(2, 7):
                r_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        sn += 1
        if isinstance(details.get('total_mark'), (int, float)):
            grand_total_marks += details.get('total_mark')
            
    # --- SUMMARY SECTION ---
    doc.add_paragraph()
    overall = terminal_data.get('overall', {})
    summary_text = (
        f"Grand Total: {grand_total_marks}         "
        f"Grade Point Average (GPA): {overall.get('gpa', 'N/A')}         "
    )
    p = doc.add_paragraph()
    p.add_run(summary_text).bold = True

    # --- FOOTER ---
    generate_footer(doc)
    add_date_of_issue(doc)

    # --- SAVE ---
    doc.save(output_filename)
    print(f"✅ Successfully created marksheet: {output_filename}")
