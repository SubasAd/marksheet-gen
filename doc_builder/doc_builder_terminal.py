
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
from .common_functions import *
from config import COMMENTS

def create_marksheet_docx(student_data, terminal_key, school_info, output_filename):
    """
    Generates a .docx marksheet for a specific terminal.
    
    Args:
        student_data (dict): The structured dictionary for one student.
        terminal_key (str): The key for the terminal (e.g., 'first_term', 'final_ledger').
        school_info (dict): A dictionary with school details.
        output_filename (str): The name of the Word file to create.
    """
    doc = docx.Document()
    
    # --- 1. HEADER SECTION ---
    # Add logo if it exists
    add_heading_component(student_data, terminal_key, school_info, doc)
    # --- 3. MARKS TABLE SECTION ---
    terminal_data = student_data[terminal_key]
    subjects = terminal_data.get('subjects', {})
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = ['S.N.', 'Subjects', 'Full Marks', 'Credit Hour', 'Marks Obtained', 'Grade', 'GP', 'Remarks']
    for i, header_text in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header_text)
        run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    total_fm = 0
    total_obtained = 0
    
    for i, (subject_name, details) in enumerate(subjects.items(), 1):
        row_cells = table.add_row().cells
        
        # Determine which marks to show based on the terminal
        if terminal_key in ['first_term', 'second_term']:
            fm = details.get('full_marks', 'N/A')
            ch = details.get('credit_hour', 'N/A')
            obtained = details.get('mark', 'N/A')
            grade = details.get('grade', 'N/A')
            gp = details.get('gp', 'N/A')
        # Add to totals for summary row
        if isinstance(fm, (int, float)): total_fm += fm
        if isinstance(obtained, (int, float)): total_obtained += obtained
        
        row_cells[0].text = str(i)
        row_cells[1].text = subject_name
        row_cells[2].text = str(fm)
        row_cells[3].text = str(ch)
        row_cells[4].text = str(obtained)
        row_cells[5].text = str(grade)
        row_cells[6].text = str(gp)
        row_cells[7].text =  COMMENTS[grade] # Remarks column is empty

    # --- 4. SUMMARY / TOTAL ROW ---
    doc.add_paragraph() # Spacer
    total_table = doc.add_table(rows=1, cols=4)
    total_table.style = 'Table Grid'
    total_cells = total_table.rows[0].cells
    
    gpa = terminal_data.get('overall', {}).get('gpa', 'N/A')
    
    total_cells[0].text = "Total"
    total_cells[0].paragraphs[0].runs[0].bold = True
    total_cells[1].text = str(total_fm)
    total_cells[2].text = str(total_obtained)
    total_cells[3].text = f"GPA: {gpa}"
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # --- 5. FOOTER SECTION ---
    generate_footer(doc)

    add_date_of_issue(doc)
    
    # --- 6. SAVE THE DOCUMENT ---
    doc.save(output_filename)
    print(f"Successfully created marksheet: {output_filename}")



