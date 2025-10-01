from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

def extract_all_student_info(student_info):
    name = student_info.get('name', ".................................")
    dob = student_info.get('dob', ".................................")
    parent = student_info.get('fathers_name',student_info.get( 'mothers_name',student_info.get('parents_name',".................................")))
    address = student_info.get('address',".................................")
    if isinstance(dob, date):
        dob = dob.strftime('%Y-%m-%d') # Format date without time

    roll_no = student_info.get('roll_number', '.......')
    return name,dob,parent,address,roll_no

# --- HELPER FUNCTIONS TO BUILD THE DOCUMENT ---
def add_logo(school_info, doc):
    try:
        # Add logo in a new paragraph at the top left, with no effect on following text
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(school_info['logo_path'], width=Cm(2.5))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except FileNotFoundError:
        print(f"Warning: Logo file not found at {school_info['logo_path']}")

def add_centered_bold_heading(doc, text, size):
    """Adds a centered, bold paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def generate_footer(doc):
    footer_table = doc.add_table(rows=1, cols=3)
    footer_cells = footer_table.rows[0].cells
    footer_cells[0].text = ".................................\nPREPARED BY"
    footer_cells[1].text = ".................................\nCHECKED BY"
    footer_cells[2].text = ".................................\nHEAD TEACHER"
    
    # Remove borders
    for cell in footer_cells:
        for key, value in cell._tc.get_or_add_tcPr().items():
            if key.endswith('Borders'):
                for border_key in value.keys():
                    value.set(border_key, 'nil')

def add_date_of_issue(doc):
    doc.add_paragraph(f"DATE OF ISSUE: {date.today().strftime('%d-%m-%Y')}")
def add_student_details(doc, student_info, student_grade, exam_date, exam_date_ad):
    """Creates the table for student details with proper formatting."""
    name, dob, parent, address, roll_no =  extract_all_student_info(student_info=student_info)
    p = doc.add_paragraph(f"""
            THE MARKS/GRADE(S) SECURED BY {name} DATE OF BIRTH  {dob}
            SYMBOL NO. {roll_no} GRADE {student_grade}
            SON/DAUGHTER OF {parent}  RESIDENT OF  {address}
            IN THE EXAMINATION CONDUCTED IN {exam_date} B.S. ({exam_date_ad} A.D.) ARE GIVEN  BELOW. \n
""")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading_component(student_data, terminal_key, school_info, doc):
    add_logo(school_info, doc)
    add_centered_bold_heading(doc, school_info['name'].upper(), 24)
    add_centered_bold_heading(doc, f"{school_info["address"].upper()}",20)
    exam_name = terminal_key.replace('_', ' ').title() + ' EXAMINATION'
    add_centered_bold_heading(doc, f"{exam_name.upper()}", 14)
    add_centered_bold_heading(doc, "GRADE SHEET", 12)

    # --- 2. STUDENT DETAILS SECTION ---
    add_student_details(doc, student_data['student_info'], school_info['grade'],school_info["exam_year_bs"], school_info["exam_year_ad"])
